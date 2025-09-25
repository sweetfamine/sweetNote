import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from typing import Any
import os
import json
from typing import cast
from openpyxl.worksheet.worksheet import Worksheet

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors


class ExportWindow(ctk.CTkToplevel):
    def __init__(self, master, manager, config, selected_ids=None):
        super().__init__(master)
        self.title("Export")
        self.geometry("400x280")

        self.transient(master)
        self.grab_set()
        self.focus_force()
        self.lift()

        self.manager = manager
        self.config = config
        self.selected_ids = selected_ids

        frame = ctk.CTkFrame(self)
        frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Auswahl: Scope
        ctk.CTkLabel(frame, text="Welche Kunden exportieren?").pack(anchor="w", pady=(0, 5))
        scope_values = ["Alle Kunden", "Nur ausgewählte"]
        default_scope = "Nur ausgewählte" if selected_ids else "Alle Kunden"
        self.scope_var = tk.StringVar(value=default_scope)
        self.scope_menu = ctk.CTkOptionMenu(frame, values=scope_values, variable=self.scope_var)
        self.scope_menu.pack(fill="x", pady=(0, 15))

        # Auswahl: Format
        ctk.CTkLabel(frame, text="Export-Format:").pack(anchor="w", pady=(0, 5))
        format_values = ["Excel (.xlsx)", "CSV (.csv)", "PDF (.pdf)", "Word (.docx)", "JSON (.json)"]
        self.format_var = tk.StringVar(value=format_values[0])
        self.format_menu = ctk.CTkOptionMenu(frame, values=format_values, variable=self.format_var)
        self.format_menu.pack(fill="x", pady=(0, 15))

        # Buttons
        btn_frame = ctk.CTkFrame(frame)
        btn_frame.pack(fill="x", pady=15)
        ctk.CTkButton(btn_frame, text="Exportieren", command=self.export).pack(side="left", expand=True, padx=10)
        ctk.CTkButton(btn_frame, text="Abbrechen", command=self.destroy).pack(side="left", expand=True, padx=10)

    def export(self) -> None:
        fmt_label = self.format_var.get()
        scope_label = self.scope_var.get()

        fmt_map = {
            "Excel (.xlsx)": "xlsx",
            "CSV (.csv)": "csv",
            "PDF (.pdf)": "pdf",
            "Word (.docx)": "docx",
            "JSON (.json)": "json",
        }
        scope = "selected" if scope_label.startswith("Nur") else "all"
        fmt = fmt_map[fmt_label]

        if scope == "selected" and self.selected_ids:
            data: list[Any] = [self.manager.get_customer_by_id(cid) for cid in self.selected_ids]
        else:
            data: list[Any] = self.manager.get_all_customers() or []

        filetypes = {
            "xlsx": ("Excel-Datei", "*.xlsx"),
            "csv": ("CSV-Datei", "*.csv"),
            "pdf": ("PDF-Datei", "*.pdf"),
            "docx": ("Word-Datei", "*.docx"),
            "json": ("JSON-Datei", "*.json"),
        }

        path = filedialog.asksaveasfilename(
            defaultextension=f".{fmt}",
            filetypes=[filetypes[fmt]],
        )
        if not path:
            return

        try:
            if fmt == "xlsx":
                self.export_excel(path, data)
            elif fmt == "csv":
                self.export_csv(path, data)
            elif fmt == "pdf":
                self.export_pdf(path, data)
            elif fmt == "docx":
                self.export_docx(path, data)
            elif fmt == "json":
                self.export_json(path, data)
            messagebox.showinfo("Export", f"Export erfolgreich:\n{path}")
            self.destroy()
        except Exception as e:
            messagebox.showerror("Fehler", f"Export fehlgeschlagen:\n{e}")

    def export_excel(self, path: str, data: list[Any]) -> None:
        from openpyxl import Workbook
        wb = Workbook()
        ws = cast(Worksheet, wb.active)

        if not data:
            ws.append(["Keine Daten"])
        else:
            ws.append(list(data[0].__dict__.keys()))
            for c in data:
                ws.append([getattr(c, k) for k in c.__dict__.keys()])
        wb.save(path)

    def export_csv(self, path: str, data: list[Any]) -> None:
        import csv
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if not data:
                writer.writerow(["Keine Daten"])
            else:
                writer.writerow(list(data[0].__dict__.keys()))
                for c in data:
                    writer.writerow([getattr(c, k) for k in c.__dict__.keys()])

    def export_pdf(self, path: str, data: list[Any]) -> None:
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        doc = SimpleDocTemplate(path, pagesize=A4)
        if not data:
            table_data = [["Keine Daten"]]
        else:
            table_data = [list(data[0].__dict__.keys())]
            for c in data:
                table_data.append([str(getattr(c, k)) for k in c.__dict__.keys()])
        table = Table(table_data)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]))
        doc.build([table])

    def export_docx(self, path: str, data: list[Any]) -> None:
        from docx import Document
        doc = Document()
        if not data:
            doc.add_paragraph("Keine Daten")
        else:
            table = doc.add_table(rows=1, cols=len(data[0].__dict__.keys()))
            hdr_cells = table.rows[0].cells
            for i, key in enumerate(data[0].__dict__.keys()):
                hdr_cells[i].text = key
            for c in data:
                row_cells = table.add_row().cells
                for i, key in enumerate(c.__dict__.keys()):
                    row_cells[i].text = str(getattr(c, key))
        doc.save(path)

    def export_json(self, path: str, data: list[Any]) -> None:
        with open(path, "w", encoding="utf-8") as f:
            if not data:
                json.dump({"info": "Keine Daten"}, f, ensure_ascii=False, indent=2)
            else:
                json.dump([c.__dict__ for c in data], f, ensure_ascii=False, indent=2)